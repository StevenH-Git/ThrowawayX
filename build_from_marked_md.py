#!/usr/bin/env python3
# pip install python-docx
#python build_from_marked_md.py --md steps.md --template Templ.docx --out Output.docx
#Anchor: ###TECHNICAL_STEPS_MD###


import argparse, re, sys
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ANCHOR = "###TECHNICAL_STEPS_MD###"

LINE_RE = re.compile(r'^(#+)\s*([+-])\s*(.*\S)?\s*$')
SPLIT_NOTE_RE = re.compile(r'\s*>>>\s*', re.IGNORECASE)

def parse_marked_lines(path):
    rows = []
    with open(path, 'r', encoding='utf-8') as f:
        for raw in f:
            s = raw.rstrip()
            if not s: continue
            m = LINE_RE.match(s)
            if not m: continue
            hashes, pm, rest = m.group(1), m.group(2), (m.group(3) or "").strip()
            h = len(hashes)
            if h == 6 and pm == '+':  # Lab Step -> skip
                continue
            if   h == 4 and pm == '+': typ = "START SECTION"
            elif h == 4 and pm == '-': typ = "END SECTION"
            elif h == 3 and pm == '+': typ = "START MAJOR"
            elif h == 3 and pm == '-': typ = "END MAJOR"
            elif h == 1 and pm == '+': typ = "STEP"
            elif h == 2 and pm == '+': typ = "NOTE"
            else: continue
            step, note = rest, ""
            if rest:
                parts = SPLIT_NOTE_RE.split(rest, maxsplit=1)
                step = parts[0].strip()
                if len(parts) > 1: note = parts[1].strip()
            rows.append((typ, step, note))
    return rows

def find_anchor_paragraph(doc, token):
    for p in doc.paragraphs:
        if token in p.text:
            return p
    return None

def insert_paragraph_after(doc, ref_el, text="", style=None):
    p = doc.add_paragraph(text)
    if style: p.style = style
    ref_el.addnext(p._element)
    return p._element

def insert_table_after(doc, ref_el, rows=1, cols=4):
    tbl = doc.add_table(rows=rows, cols=cols)
    ref_el.addnext(tbl._tbl)
    return tbl, tbl._tbl

def content_width(doc):
    sec = doc.sections[0]
    return sec.page_width - sec.left_margin - sec.right_margin  # twips

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.get_or_add_tblPr()
    # remove existing borders
    for e in list(tblPr):
        if e.tag == qn('w:tblBorders'):
            tblPr.remove(e)
    borders = OxmlElement('w:tblBorders')
    def edge(tag):
        e = OxmlElement(tag)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '8')   # 0.5pt
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), 'auto')
        return e
    for tag in ('w:top','w:left','w:bottom','w:right','w:insideH','w:insideV'):
        borders.append(edge(tag))
    tblPr.append(borders)

def zero_par_spacing(cell):
    for p in cell.paragraphs:
        pf = p.paragraph_format
        pf.space_before = 0
        pf.space_after = 0

def apply_vba_like_widths(doc, table):
    # fixed w1 ≈ 0.9", w4 ≈ 0.5" (fall back to 0.3" if needed), body = remainder
    cw = content_width(doc)  # twips
    w1 = Inches(0.9)
    w4 = Inches(0.5)
    min_body = Inches(1.0)
    if cw < (w1 + w4 + min_body):
        w4 = Inches(0.3)
        if cw < (w1 + w4 + min_body):
            w1 = Inches(0.7)
    body = max(Inches(0.5), cw - w1 - w4)
    w2 = body * 0.55
    w3 = body - w2
    table.autofit = False
    table.allow_autofit = False  # some builds expose this alias
    cols = table.columns
    # set per-column via per-row cells (python-docx quirk)
    for r in table.rows:
        r.cells[0].width = w1
        r.cells[1].width = w2
        r.cells[2].width = w3
        r.cells[3].width = w4
        for c in r.cells:
            zero_par_spacing(c)

def build_after_anchor(doc, anchor_p, rows):
    # headings similar to VBA: bold, Heading 2/3
    for sname in ("Heading 2","Heading 3"):
        try:
            st = doc.styles[sname]; st.font.size = Pt(12 if sname=="Heading 2" else 11); st.font.bold = True
        except Exception:
            pass

    sec_idx = maj_idx = stp_idx = 0
    current_tbl = None
    first_row_written = False
    ref_el = anchor_p._element

    for (typ, step, note) in rows:
        if typ == "START SECTION":
            current_tbl = None; maj_idx = 0; stp_idx = 0; sec_idx += 1
            ref_el = insert_paragraph_after(doc, ref_el, step, style="Heading 2")
            ref_el = insert_paragraph_after(doc, ref_el, "")
        elif typ == "END SECTION":
            ref_el = insert_paragraph_after(doc, ref_el, "")
        elif typ == "START MAJOR":
            current_tbl = None; maj_idx += 1; stp_idx = 0
            ref_el = insert_paragraph_after(doc, ref_el, step if step else "Major", style="Heading 3")
            current_tbl, ref_el = insert_table_after(doc, ref_el, rows=1, cols=4)
            set_table_borders(current_tbl)
            apply_vba_like_widths(doc, current_tbl)
            first_row_written = False
        elif typ == "END MAJOR":
            current_tbl = None
            ref_el = insert_paragraph_after(doc, ref_el, "")
        elif typ == "STEP":
            if current_tbl is None: 
                continue
            stp_idx += 1
            idx = f"{sec_idx}.{maj_idx}.{stp_idx}"
            if not first_row_written:
                cells = current_tbl.rows[0].cells
                cells[0].text = idx; cells[1].text = step; cells[2].text = note
                for c in cells: zero_par_spacing(c)
                first_row_written = True
            else:
                rc = current_tbl.add_row().cells
                rc[0].text = idx; rc[1].text = step; rc[2].text = note
                for c in rc: zero_par_spacing(c)
            apply_vba_like_widths(doc, current_tbl)
        elif typ == "NOTE":
            # ignored per spec
            pass

    # remove the anchor paragraph
    a = anchor_p._element
    a.getparent().remove(a)
    a._p = a._element = None

def main():
    ap = argparse.ArgumentParser(description="Insert marked MD steps after anchor in DOCX (VBA-matched table formatting)")
    ap.add_argument("--md", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--template", required=True)
    args = ap.parse_args()

    rows = parse_marked_lines(args.md)
    if not rows:
        print("No valid lines parsed.", file=sys.stderr); sys.exit(1)

    doc = Document(args.template)
    anchor_p = find_anchor_paragraph(doc, ANCHOR)
    if anchor_p is None:
        print(f"Anchor not found: {ANCHOR}", file=sys.stderr); sys.exit(2)

    build_after_anchor(doc, anchor_p, rows)
    doc.save(args.out)

if __name__ == "__main__":
    main()
